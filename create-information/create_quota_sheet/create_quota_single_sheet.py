import psycopg2
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_TAB_LEADER
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import os
import sys
import codecs
from datetime import datetime

def writeOrderNumberDocument(sOrderNumber):
	global conn
	global lRow
	
	print ("Creating document for quota " + sOrderNumber)
	sFilename = "C:\\projects\\create_quota_sheet\\word\\" + sOrderNumber + ".docx"
	document = Document()
	sections = document.sections
	margin = 1.5
	for section in sections:
		section.top_margin = Cm(margin)
		section.bottom_margin = Cm(margin)
		section.left_margin = Cm(margin)
		section.right_margin = Cm(margin)
	
		
	# Create styles
	my_styles = document.styles
	p_style = my_styles.add_style('Normal in table', WD_STYLE_TYPE.PARAGRAPH)
	p_style.base_style = my_styles['Normal']
	p_style.font.name = "Calibri"
	p_style.paragraph_format.space_before = Pt(2)
	p_style.paragraph_format.space_after = Pt(2)
	
	p_style2 = my_styles.add_style('Small in table', WD_STYLE_TYPE.PARAGRAPH)
	p_style2.base_style = my_styles['Normal in table']
	p_style2.font.size = Pt(9)
	
	h1_style = my_styles["Heading 1"]
	h1_style.paragraph_format.space_before = Pt(12)
	
	
	# The main title
	document.add_heading("Quota order number " + sOrderNumber, level = 0)
	
	# Background information
	document.add_heading("About this quota", level = 1)
	p = document.add_paragraph('The table below contains background information provided by Defra.')
	p.style = p_style2
	sSQL = """SELECT * FROM ml.quota_extension WHERE quota_order_number_id = '""" + sOrderNumber + """'"""
	
	cur = conn.cursor()
	cur.execute(sSQL)
	rw = cur.fetchone()
	table = document.add_table(rows = 10, cols = 2)
	table.style = "Table Grid"

	table.rows[0].cells[0].text = "Defra description"
	table.rows[1].cells[0].text = "Country"
	table.rows[2].cells[0].text = "Application period"
	table.rows[3].cells[0].text = "Overall quota volume"
	table.rows[4].cells[0].text = "Unit"
	table.rows[5].cells[0].text = "Method"
	table.rows[6].cells[0].text = "Type"
	table.rows[7].cells[0].text = "Mechanism"
	table.rows[8].cells[0].text = "Sector"
	table.rows[9].cells[0].text = "Comments"

	if not(rw is None):
		table.rows[0].cells[1].text = m_str(rw[1])
		table.rows[1].cells[1].text = m_str(rw[3])
		table.rows[2].cells[1].text = m_str(rw[4])
		table.rows[3].cells[1].text = "{:,}".format(rw[5])
		table.rows[4].cells[1].text = m_str(rw[6])
		table.rows[5].cells[1].text = m_str(rw[7])
		table.rows[6].cells[1].text = m_str(rw[8])
		table.rows[7].cells[1].text = m_str(rw[9])
		table.rows[8].cells[1].text = m_str(rw[10])
		table.rows[9].cells[1].text = m_str(rw[11])

	for i in range(10):
		r = table.rows[i].cells
		r[0].paragraphs[0].style = p_style2
		r[1].paragraphs[0].style = p_style2
	

	set_col_widths(table, (Cm(4.6), Cm(14)))
		
	# Regulations
	document.add_heading("Regulation(s)", level = 1)
	p = document.add_paragraph('The table below lists the EU regulation(s) that has/have given rise to this quota and its associated measures.')
	p.style = p_style2
	sSQL = """SELECT DISTINCT LEFT(m.measure_generating_regulation_id, 7) as regulation_id, r.title
	FROM measures m , ml.ml_regulation_titles r 
	WHERE LEFT(m.measure_generating_regulation_id, 7) = r.regulation_id
	AND m.ordernumber IS NOT NULL
	AND m.validity_start_date > (CURRENT_DATE - 365)
	AND ordernumber = '""" + sOrderNumber + """'
	ORDER BY 1;"""
	
	cur = conn.cursor()
	cur.execute(sSQL)
	rws = cur.fetchall()
	table = document.add_table(rows = cur.rowcount + 1, cols = 2)
	table.style = "Light List"
	hdr_cells = table.rows[0].cells
	set_repeat_table_header(table.rows[0])
	hdr_cells[0].text = 'EU Regulation ID'
	hdr_cells[1].text = 'Title'
	hdr_cells[0].paragraphs[0].style = p_style2
	hdr_cells[1].paragraphs[0].style = p_style2
	i = 1
	for rw in rws:
		regulation_id = rw[0]
		title         = rw[1]
		r = table.rows[i].cells
		r[0].text = regulation_id
		r[1].text = title
		r[0].paragraphs[0].style = p_style2
		r[1].paragraphs[0].style = p_style2
		i = i + 1

	set_col_widths(table, (Cm(4), Cm(14.6)))
		
	# Geography
	document.add_heading("Geography", level = 1)
	p = document.add_paragraph('The table below lists the geographies to which the quota applies.')
	p.style = p_style2
	sSQL = """SELECT DISTINCT qono.geographical_area_id, ga.description
	FROM quota_order_numbers qon, quota_order_number_origins qono, ml.ml_geographical_areas ga
	WHERE qon.quota_order_number_sid = qono.quota_order_number_sid
	AND ga.geographical_area_id = qono.geographical_area_id
	AND (qono.validity_end_date IS NULL OR qono.validity_end_date > CURRENT_DATE)
	AND qon.quota_order_number_id = '""" + sOrderNumber + """'
	ORDER BY 1;"""
	
	cur = conn.cursor()
	cur.execute(sSQL)
	rws = cur.fetchall()
	table = document.add_table(rows = cur.rowcount + 1, cols = 2)
	table.style = "Light List"
	hdr_cells = table.rows[0].cells
	set_repeat_table_header(table.rows[0])
	hdr_cells[0].text = 'Geographical area ID'
	hdr_cells[1].text = 'Description'
	hdr_cells[0].paragraphs[0].style = p_style2
	hdr_cells[1].paragraphs[0].style = p_style2
	i = 1
	for rw in rws:
		geographical_area_id = rw[0]
		description          = rw[1]
		r = table.rows[i].cells
		r[0].text = geographical_area_id
		r[1].text = description
		r[0].paragraphs[0].style = p_style2
		r[1].paragraphs[0].style = p_style2
		i = i + 1

	set_col_widths(table, (Cm(4), Cm(14.6)))
		
	
	# Geographical area exclusions
	document.add_heading("Geographical area exclusions", level = 2)
	sSQL = """
	SELECT qono.geographical_area_id, gaparent.description as origin, gachild.geographical_area_id as excluded_orgin_id, gachild.description as excluded_orgin
	FROM quota_order_numbers qon, quota_order_number_origin_exclusions qonoe, quota_order_number_origins qono, ml.ml_geographical_areas gaparent, ml.ml_geographical_areas gachild
	WHERE qonoe.quota_order_number_origin_sid = qono.quota_order_number_origin_sid
	AND qon.quota_order_number_sid = qono.quota_order_number_sid
	AND qono.geographical_area_id = gaparent.geographical_area_id
	AND qonoe.excluded_geographical_area_sid = gachild.geographical_area_sid
	AND (qono.validity_end_date IS NULL OR qono.validity_end_date > CURRENT_DATE)
	AND qon.quota_order_number_id = '""" + sOrderNumber + """' ORDER BY 1, 3;
	"""
	cur = conn.cursor()
	cur.execute(sSQL)
	rws = cur.fetchall()
	if cur.rowcount > 0:
		p = document.add_paragraph('The table below lists the exclusions from the aforementioned geographical areas.')
		p.style = p_style2
		table = document.add_table(rows = cur.rowcount + 1, cols = 2)
		table.style = "Light List"
		hdr_cells = table.rows[0].cells
		set_repeat_table_header(table.rows[0])
		hdr_cells[0].text = 'Parent area ID'
		hdr_cells[1].text = 'Excluded area ID'
		hdr_cells[0].paragraphs[0].style = p_style2
		hdr_cells[1].paragraphs[0].style = p_style2
		i = 1
		for rw in rws:
			geographical_area_id1 = rw[0]
			description1          = rw[1]
			geographical_area_id2 = rw[2]
			description2          = rw[3]
			r = table.rows[i].cells
			r[0].text = geographical_area_id1 + " - " + description1
			r[1].text = geographical_area_id2 + " - " + description2
			r[0].paragraphs[0].style = p_style2
			r[1].paragraphs[0].style = p_style2
			i = i + 1
		
		set_col_widths(table, (Cm(9.3), Cm(9.3)))
	else:
		p = document.add_paragraph('No geographical exclusions.')
		p.style = p_style2
		

	# Commodity codes
	document.add_heading("Commodity codes", level = 1)
	p = document.add_paragraph('The table below lists the commodity codes to which this quota applies.')
	p.style = p_style2
	sSQL = """
	SELECT DISTINCT m.goods_nomenclature_item_id, n.description
	FROM measures m, ml.ml_nomenclature_flat n
	WHERE n.product_line_suffix = '80'
	AND n.goods_nomenclature_item_id = m.goods_nomenclature_item_id
	AND m.ordernumber = '""" + sOrderNumber + """'
	AND (m.validity_end_date IS NULL OR m.validity_end_date > CURRENT_DATE)
	AND m.validity_start_date > (CURRENT_DATE - 365)
	ORDER BY 1;
	"""
	cur = conn.cursor()
	cur.execute(sSQL)
	rws = cur.fetchall()
	table = document.add_table(rows = cur.rowcount + 1, cols = 2)
	table.style = "Light List"
	hdr_cells = table.rows[0].cells
	set_repeat_table_header(table.rows[0])
	hdr_cells[0].text = 'Commodity code'
	hdr_cells[1].text = 'Description'
	hdr_cells[0].paragraphs[0].style = p_style2
	hdr_cells[1].paragraphs[0].style = p_style2
	i = 1
	for rw in rws:
		commcode    = rw[0]
		description = rw[1]
		r = table.rows[i].cells
		r[0].text = commcode
		r[1].text = description
		r[0].paragraphs[0].style = p_style2
		r[1].paragraphs[0].style = p_style2
		i = i + 1

	set_col_widths(table, (Cm(3.5), Cm(15.1)))
	
	# Full list of measures
	document.add_heading("Measures and duties", level = 1)
	sSQL = """
	SELECT DISTINCT m.measure_type_id, mtd.description as measure_type, m.goods_nomenclature_item_id, mc.duty_expression_id, mc.duty_amount,
	mc.monetary_unit_code, mc.measurement_unit_code, mc.measurement_unit_qualifier_code, m.ordernumber, m.measure_sid
	FROM measures m, measure_components mc, measure_type_descriptions mtd, base_regulations r
	WHERE m.measure_sid = mc.measure_sid
	AND m.measure_generating_regulation_id = r.base_regulation_id
	AND m.measure_type_id = mtd.measure_type_id
	AND m.ordernumber IS NOT NULL
	AND r.effective_end_date IS NULL
	AND m.ordernumber = '""" + sOrderNumber + """'
	AND m.stopped_flag = 'f'
	AND (m.validity_end_date > CURRENT_DATE OR m.validity_end_date IS NULL)
	ORDER BY ordernumber, goods_nomenclature_item_id, duty_expression_id;
	"""
	cur = conn.cursor()
	cur.execute(sSQL)
	if cur.rowcount == 0:
		p = document.add_paragraph('There are no current measures - please check with the policy owner (the measures may have just expired).')
		p.style = p_style2
	else:
		p = document.add_paragraph('The table below shows the measures that area applicable to goods in quota. In most cases, the duty will be 0%, however on occasion duties remain applicable. Please also be aware that on occasion, more than one measure type may apply to a single order number.')
		p.style = p_style2
		rws = cur.fetchall()
		table = document.add_table(rows = 1, cols = 3)
		table.style = "Light List"
		hdr_cells = table.rows[0].cells
		set_repeat_table_header(table.rows[0])
		hdr_cells[0].text = 'Commodity code'
		hdr_cells[1].text = 'Measure type'
		hdr_cells[2].text = 'Applicable duty'
		hdr_cells[0].paragraphs[0].style = p_style2
		hdr_cells[1].paragraphs[0].style = p_style2
		hdr_cells[2].paragraphs[0].style = p_style2
		rdList = []
		for rw in rws:
			measure_type_id                 = rw[0]
			measure_type                    = rw[1]
			goods_nomenclature_item_id      = rw[2]
			duty_expression_id              = rw[3]
			duty_amount                     = str(rw[4])
			monetary_unit_code              = m_str(rw[5])
			measurement_unit_code           = m_str(rw[6])
			measurement_unit_qualifier_code = m_str(rw[7])
			measure_sid                     = rw[9]
			dex  = getDuty(duty_expression_id, duty_amount, monetary_unit_code, measurement_unit_code, measurement_unit_qualifier_code)
			rdList.append([measure_type_id, measure_type, goods_nomenclature_item_id, dex, measure_sid, "Active"])
			
		measure_sid_old = rdList[0][4]
		for x in range(1, len(rdList) - 1):
			measure_sid = rdList[x][4]
			if measure_sid == measure_sid_old:
				rdList[x - 1][3] += rdList[x][3]
				rdList[x][5] = "Inactive"
				
			measure_sid_old = measure_sid

		for x in range(0, len(rdList) - 1):
			if rdList[x][5] == "Active":
			
				rx = table.add_row()
				rc = rx.cells

				rc[0].text = rdList[x][2]
				rc[1].text = rdList[x][0] + " [" + rdList[x][1] + "]"
				rc[2].text = rdList[x][3]
				
				rc[0].paragraphs[0].style = p_style2
				rc[1].paragraphs[0].style = p_style2
				rc[2].paragraphs[0].style = p_style2
		set_col_widths(table, (Cm(3.5), Cm(8.3), Cm(6.8)))

	###############################################################
	# Get the definition description
	document.add_heading("Quota definition description", level = 1)
	p = document.add_paragraph("""The table cell below identifies the description that the EU has applied to the quota definition periods listed below. There may be multiple descriptions, however to aid readability, only one is included. Where there are multiple descriptions, they are usually the same.""")
	p.style = p_style2
	sSQL = """SELECT DISTINCT qd.description
	FROM measures m, quota_definitions qd, quota_order_numbers qon
	WHERE m.ordernumber = qd.quota_order_number_id
	AND m.ordernumber = qon.quota_order_number_id
	AND m.validity_start_date > (CURRENT_DATE - 365)
	AND qd.validity_start_date > (CURRENT_DATE - 365)
	AND m.ordernumber = '""" + sOrderNumber + """'
	ORDER BY 1 DESC LIMIT 1;"""
	
	cur = conn.cursor()
	cur.execute(sSQL)
	rws = cur.fetchall()

	table = document.add_table(rows = 1, cols = 1)
	table.style = "Table Grid"
	for rw in rws:
		description = rw[0]
		r = table.rows[0].cells
		r[0].text = description
		r[0].paragraphs[0].style = p_style2

	

	###############################################################
	# Get the definition periods
	document.add_heading("Quota definition periods", level = 1)
	p = document.add_paragraph('The table below lists the quota definition periods applicable to this quota.')
	p.style = p_style2
	sSQL = """SELECT DISTINCT qd.validity_start_date as defintion_start_date, qd.validity_end_date as defintion_end_date,
	qd.initial_volume, qd.measurement_unit_code, qd.maximum_precision, qd.critical_state, qd.critical_threshold, qd.monetary_unit_code, qd.measurement_unit_qualifier_code
	FROM measures m, quota_definitions qd, quota_order_numbers qon
	WHERE m.ordernumber = qd.quota_order_number_id
	AND m.ordernumber = qon.quota_order_number_id
	AND m.validity_start_date > (CURRENT_DATE - 365)
	AND qd.validity_start_date > (CURRENT_DATE - 365)
	AND m.ordernumber = '""" + sOrderNumber + """'
	ORDER BY 1, 2;"""
	
	cur = conn.cursor()
	cur.execute(sSQL)
	rows_definitions = cur.fetchall()

	table = document.add_table(rows = cur.rowcount + 1, cols = 7)
	table.style = "Light List"
	hdr_cells = table.rows[0].cells
	set_repeat_table_header(table.rows[0])
	hdr_cells[0].text = 'Dates'
	hdr_cells[1].text = 'Period type'
	hdr_cells[2].text = 'Volume'
	hdr_cells[3].text = 'Unit'
	hdr_cells[4].text = 'Monetary unit'
	hdr_cells[5].text = 'Critical state'
	hdr_cells[6].text = 'Critical threshold'
	for x in range(7):
		hdr_cells[x].paragraphs[0].style = p_style2

	i = 0
	for rw in rows_definitions:
		i = i + 1
		definition_start_date           = fmtDate(rw[0])
		definition_end_date             = fmtDate(rw[1])
		initial_volume                  = "{:,}".format(rw[2])
		measurement_unit_code           = m_str(rw[3])
		critical_state                  = rw[5]
		critical_threshold              = m_str(rw[6])
		monetary_unit_code              = m_str(rw[7])
		measurement_unit_qualifier_code = m_str(rw[8])
		
		r = table.rows[i].cells
		r[0].text = definition_start_date + " to " + definition_end_date
		r[1].text = describe_period(rw[0], rw[1])
		r[2].text = str(initial_volume)
		r[3].text = measurement_unit_code + " " + measurement_unit_qualifier_code
		r[4].text = monetary_unit_code
		r[5].text = critical_state
		r[6].text = critical_threshold

		for x in range(7):
			r[x].paragraphs[0].style = p_style2

	set_col_widths(table, (Cm(3.5), Cm(2.7), Cm(2.9), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5)))

	###############################################################
	# Get the quota associations
	document.add_heading("Quota associations", level = 1)
	sSQL = """SELECT relation_type, coefficient, qdm.quota_order_number_id as main_id, qds.quota_order_number_id as sub_id
	FROM quota_associations qa, quota_definitions qdm, quota_definitions qds
	WHERE qa.main_quota_definition_sid = qdm.quota_definition_sid
	AND qa.sub_quota_definition_sid = qds.quota_definition_sid
	AND (qdm.quota_order_number_id = '""" + sOrderNumber + """'
	OR qds.quota_order_number_id = '""" + sOrderNumber + """')
	AND qdm.validity_end_date > CURRENT_DATE
	AND qds.validity_end_date > CURRENT_DATE
	"""
	
	cur = conn.cursor()
	cur.execute(sSQL)
	rws = cur.fetchall()
	if cur.rowcount > 0:
		p = document.add_paragraph('The table below lists the quota associations applicable to this quota.')
		p.style = p_style2
		table = document.add_table(rows = cur.rowcount + 1, cols = 4)
		table.style = "Light List"
		hdr_cells = table.rows[0].cells
		set_repeat_table_header(table.rows[0])
		hdr_cells[0].text = 'Main order number'
		hdr_cells[1].text = 'Sub order number'
		hdr_cells[2].text = 'Relation type'
		hdr_cells[3].text = 'Coefficient'
		for x in range(4):
			hdr_cells[x].paragraphs[0].style = p_style

		i = 0
		for rw in rws:
			i = i + 1
			relation_type = rw[0]
			coefficient   = "{:,}".format(rw[1])
			main_id       = rw[2]
			sub_id        = rw[3]
			
			r = table.rows[i].cells
			r[0].text = main_id
			r[1].text = sub_id
			r[2].text = relation_type
			r[3].text = coefficient
			
			for x in range(4):
				r[x].paragraphs[0].style = p_style
	else:
		p = document.add_paragraph('There are no quota associations applicable to this quota.')
		p.style = p_style2
		
	###############################################################
	# Get the suspension periods
	document.add_heading("Suspension periods", level = 1)
	p = document.add_paragraph('If you know of any quota suspension periods that are due to take effect, please enter the dates in the table below')
	p.style = p_style2
	table = document.add_table(rows = 3, cols = 3)
	table.style = "Light List"
	hdr_cells = table.rows[0].cells
	set_repeat_table_header(table.rows[0])
	hdr_cells[0].text = 'Suspension start'
	hdr_cells[1].text = 'Suspension end'
	hdr_cells[2].text = 'Reason'
	for x in range(3):
		for y in range(3):
			rw = table.rows[y].cells
			rw[x].paragraphs[0].style = p_style2

	document.save(sFilename)
	#sys.exit()


def m_str(s):
	if s is None:
		return ""
	else:
		return str(s)
	
def describe_period(definition_start_date, definition_end_date):
	period_length = definition_end_date - definition_start_date
	if period_length == 365:
		sOut = "Annual"
	elif period_length in (28, 30, 31):
		sOut = "One month"
	elif period_length in (88, 89, 90, 91, 92):
		sOut = "Quarter"
	else:
		iDayStart = definition_start_date.day
		iDayEnd = definition_end_date.day
		if iDayStart == 1 and iDayEnd >= 28:
			iMonthStart = definition_start_date.month
			iMonthEnd = definition_end_date.month

			iYearStart = definition_start_date.year
			iYearEnd = definition_end_date.year

			iMonths = (iMonthEnd - iMonthStart) + (12 * (iYearEnd - iYearStart)) + 1
			sOut = str(iMonths) + " months"
		else:
			sOut = "Custom period"

	return sOut

def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
	
def fmtDate(d):
	try:
		d = datetime.strftime(d, '%d-%m-%y')
	except:
		d = ""
	return d

def getDuty(duty_expression_id, duty_amount, monetary_unit_code, measurement_unit_code, measurement_unit_qualifier_code):
	if duty_expression_id == "01": # % or amount
		return duty_amount + "%"
	elif duty_expression_id in ("04", "19"): # + % or amount
		s = " + " + duty_amount
		if monetary_unit_code != "":
			s = s + " " + monetary_unit_code + " / " + getMeasurementUnit(measurement_unit_code)
			if measurement_unit_qualifier_code != "":
				s = s + " " + measurement_unit_qualifier_code
		return s
	elif duty_expression_id == "14": # + reduced agricultural component
		return " + reduced agricultural component"
	elif duty_expression_id in ("17", "35"): # Maximum
		s = " up to a maximum of " + duty_amount
		if monetary_unit_code != "":
			s = s + " " + monetary_unit_code + " / " + getMeasurementUnit(measurement_unit_code)
			if measurement_unit_qualifier_code != "":
				s = s + " " + measurement_unit_qualifier_code
		return s
	elif duty_expression_id == "25": # + reduced additional duty on sugar
		return " + reduced additional duty on sugar"
	elif duty_expression_id == "29": # + reduced additional duty on flour
		return " + reduced additional duty on flour"

	return duty_expression_id
	

def getMeasurementUnit(s):
	if s == "ASV":
		return "% vol/hl" # 3302101000
	if s == "NAR":
		return "p/st"
	elif s == "CCT":
		return "ct/l"
	elif s == "CEN":
		return "100 p/st"
	elif s == "CTM":
		return "c/k"
	elif s == "DTN":
		return "100 kg"
	elif s == "GFI":
		return "gi F/S"
	elif s == "GRM":
		return "g"
	elif s == "HLT":
		return "hl" # 2209009100
	elif s == "HMT":
		return "100 m" # 3706909900
	elif s == "KGM":
		return "kg"
	elif s == "KLT":
		return "1,000 l"
	elif s == "KMA":
		return "kg met.am."
	elif s == "KNI":
		return "kg N"
	elif s == "KNS":
		return "kg H2O2"
	elif s == "KPH":
		return "kg KOH"
	elif s == "KPO":
		return "kg K2O"
	elif s == "KPP":
		return "kg P2O5"
	elif s == "KSD":
		return "kg 90 % sdt"
	elif s == "KSH":
		return "kg NaOH"
	elif s == "KUR":
		return "kg U"
	elif s == "LPA":
		return "l alc. 100%"
	elif s == "LTR":
		return "l"
	elif s == "MIL":
		return "1,000 p/st"
	elif s == "MTK":
		return "m2"
	elif s == "MTQ":
		return "m3"
	elif s == "MTR":
		return "m"
	elif s == "MWH":
		return "1,000 kWh"
	elif s == "NCL":
		return "ce/el"
	elif s == "NPR":
		return "pa"
	elif s == "TJO":
		return "TJ"
	elif s == "TNE":
		return "t" # 1005900020
		# return "1000 kg" # 1005900020
	else:
		return s
	
conn = psycopg2.connect("dbname=trade_tariff_1809 user=postgres password=zanzibar")

sSQL = """SELECT DISTINCT m.ordernumber FROM measures m WHERE m.validity_start_date > (CURRENT_DATE - 365) ORDER BY 1;"""
cur = conn.cursor()
cur.execute(sSQL)
rows_ordernumbers = cur.fetchall()
for rw in rows_ordernumbers:
	sOrderNumber = rw[0]
	#if sOrderNumber > "090101":
	writeOrderNumberDocument(sOrderNumber)

conn.close()
