import psycopg2
import mistune
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_TAB_LEADER
from __future__ import with_statement
from contextlib import closing
from zipfile import ZipFile, ZIP_DEFLATED
import os
import sys

conn = psycopg2.connect("dbname=trade_tariff_1809 user=postgres password" + self.p)

def zipdir(basedir, archivename):
	assert os.path.isdir(basedir)
	with closing(ZipFile(archivename, "w", ZIP_DEFLATED)) as z:
		for root, dirs, files in os.walk(basedir):
			#NOTE: ignore empty directories
			for fn in files:
				absfn = os.path.join(root, fn)
				zfn = absfn[len(basedir)+len(os.sep):] #XXX: relative path
				z.write(absfn, zfn)

def formatCommodityCode(s):
	s2 = s[0:4] + ' ' + s[4:6] + ' ' + s[6:8] + ' ' + s[8:10]
	return (s2)


def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row
	
###############################################################
# Define the parameters
try:
	sChapter = str(sys.argv[1])
	sChapter = sChapter.zfill(2)
	sDocumentType = str(sys.argv[2])
	if (sDocumentType != "classification" and sDocumentType != "schedule"):
		sDocumentType = "classification"
except:
	sChapter = "01"	
	sDocumentType = "classification"
	
print (sDocumentType)

if (sDocumentType == "classification"):
	sTitle = "UK Goods Classification"
	intColumnCount = 2
	sOutTemplate = "output_classification_XX.docx"
else:
	sTitle = "UK Goods Schedule"
	intColumnCount = 3
	sOutTemplate = "output_schedule_XX.docx"


###############################################################
# Step 1 - get the section header
# Relevant to both the classification and the schedule
sSQL = """SELECT s.numeral, s.title, cs.section_id
FROM goods_nomenclatures gn, chapters_sections cs, sections s
WHERE gn.goods_nomenclature_sid = cs.goods_nomenclature_sid
AND s.id = cs.section_id
AND gn.goods_nomenclature_item_id = '""" + sChapter + """00000000'"""
cur = conn.cursor()
cur.execute(sSQL)
row = cur.fetchone()
sSectionNumeral = row[0]
sSectionTitle = row[1]
#print(sSectionTitle)

sSectionNumeralAndTitle = "Section " + sSectionNumeral + " " + sSectionTitle

###############################################################
# Step 2 - Get the chapter details
# Relevant to both the classification and the schedule
sSQL = "SELECT description FROM ml.chapters WHERE chapter = '" + sChapter + "'"
cur = conn.cursor()
cur.execute(sSQL)
row = cur.fetchone()
sChapterDescription = row[0]
sChapter2 = "Chapter " + sChapter + " " + sChapterDescription

###############################################################
# Step 3 - Get the chapter notes
# Relevant to both the classification and the schedule
sSQL = "SELECT content FROM chapter_notes WHERE chapter_id = '" + sChapter + "'"
cur = conn.cursor()
cur.execute(sSQL)
row = cur.fetchone()
sChapterNotes = row[0]

renderer = mistune.Renderer(escape=True, hard_wrap=True)
markdown = mistune.Markdown(renderer=renderer)
sChapterNotes2 = markdown(sChapterNotes)

file = open("chapter.md", "w")
file.write(sChapterNotes2)
file.close()
cur.close()

###############################################################
# Step 4 - Get the duties
# Relevant to just the schedule
if (sDocumentType == "schedule"):
	sSQL = """SELECT m.goods_nomenclature_item_id, m.measure_type_id, mc.duty_expression_id, mc.duty_amount, mc.monetary_unit_code,
	mc.measurement_unit_code, mc.measurement_unit_qualifier_code, m.validity_start_date, m.validity_end_date, m.geographical_area_id
	FROM measure_components mc, ml.v5_2019 m
	WHERE mc.measure_sid = m.measure_sid
	AND LEFT(m.goods_nomenclature_item_id, 2) ='""" + sChapter + """'
	/*AND m.validity_start_date < CURRENT_DATE
	AND (m.validity_end_date > CURRENT_DATE OR m.validity_end_date IS NULL)*/
	AND m.measure_type_id IN ('103', '105')
	ORDER BY m.goods_nomenclature_item_id, mc.duty_expression_id"""
	cur = conn.cursor()
	cur.execute(sSQL)
	rows_duties = cur.fetchall()

	# Do a pass through the duties table and create a full duty expression
	rdList = []
	for rd in rows_duties:
		sCommodityCode = rd[0]
		sDutyExpression = str(rd[2])
		sDutyAmount = str(rd[3])
		sMonetaryUnitCode = str(rd[4])
		sMeasurementUnitCode = str(rd[5])
		sMeasurementUnitQualifierCode = str(rd[6])
		
		sFullExpression = sDutyAmount
		if (sMonetaryUnitCode == "" or sMonetaryUnitCode == "None"):
			sFullExpression += "%"
		else:
			sFullExpression += " " + sMonetaryUnitCode
			
		if (sMeasurementUnitCode != "" and sMeasurementUnitCode != "None"):
			sFullExpression += " / " + sMeasurementUnitCode

		rdList.append([sCommodityCode, sDutyExpression, sDutyAmount, sMonetaryUnitCode, sMeasurementUnitCode, sMeasurementUnitQualifierCode, sFullExpression, "Active"])
		
	# Now, do a pass through the duties table and join up where there are multiple
	sCommodityCodeOld = ""

	for x in range(1, len(rdList) - 1):
		sCommodityCode = rdList[x][0]
		sDutyExpression = str(rdList[x][1])
		sDutyAmount = str(rdList[x][2])
		sMonetaryUnitCode = str(rdList[x][3])
		sMeasurementUnitCode = str(rdList[x][4])
		sMeasurementUnitQualifierCode = str(rdList[x][5])
		sFullExpression = str(rdList[x][6])
		
		if (sCommodityCodeOld == sCommodityCode):
			rdList[x - 1][6] += " + " + rdList[x][6]
			rdList[x][7] = "Inactive"
			#print ("Joining")
		
		sCommodityCodeOld = sCommodityCode

# Step 4b - Get the supplementary units
sSQL = """SELECT * FROM ml.v5_2019 WHERE measure_type_id IN ('109', '110') AND goods_nomenclature_item_id LIKE '""" + sChapter + """%'"""
curSupplementary = conn.cursor()
curSupplementary.execute(sSQL)
rowsSupplementary = curSupplementary.fetchall()


###############################################################
# Step 5 - Get the table of classifications
# Relevant to just the schedule
sSQL = """SELECT gn.goods_nomenclature_item_id, gn.producline_suffix, gn.validity_start_date, gn.validity_end_date, 
gnd.productline_suffix, gnd.description, gni.number_indents, 
gndp.goods_nomenclature_description_period_sid, gndp.validity_start_date, gndp.validity_end_date,
concat(repeat (' -', gni.number_indents), ' ', gnd.description) "nice_description",
left (gn.goods_nomenclature_item_id, 2) "chapter", CASE gn.producline_suffix WHEN '80' THEN 1 ELSE 0 END "leaf" 
FROM goods_nomenclatures gn
JOIN goods_nomenclature_descriptions gnd ON gnd.goods_nomenclature_sid = gn.goods_nomenclature_sid
JOIN goods_nomenclature_description_periods gndp ON gndp.goods_nomenclature_description_period_sid = gnd.goods_nomenclature_description_period_sid
JOIN goods_nomenclature_indents gni ON gni.goods_nomenclature_sid = gn.goods_nomenclature_sid
WHERE (gn.validity_end_date IS NULL OR gn.validity_end_date >= CURRENT_DATE)
AND gndp.goods_nomenclature_description_period_sid IN
(SELECT MAX (goods_nomenclature_description_period_sid)
FROM goods_nomenclature_description_periods gndp2
WHERE gndp2.goods_nomenclature_sid = gnd.goods_nomenclature_sid
AND gndp2.validity_start_date < CURRENT_DATE
UNION
SELECT goods_nomenclature_description_period_sid
FROM goods_nomenclature_description_periods gndp3
WHERE gndp3.goods_nomenclature_sid = gnd.goods_nomenclature_sid
AND gndp3.validity_start_date >= CURRENT_DATE)
AND gni.goods_nomenclature_indent_sid IN
(SELECT MAX (goods_nomenclature_indent_sid)
FROM goods_nomenclature_indents gni2
WHERE gni2.goods_nomenclature_sid = gn.goods_nomenclature_sid
AND gni2.validity_start_date < CURRENT_DATE
UNION
SELECT goods_nomenclature_indent_sid
FROM goods_nomenclature_indents gni3
WHERE gni3.goods_nomenclature_sid = gn.goods_nomenclature_sid
AND gni3.validity_start_date >= CURRENT_DATE) 
AND left (gn.goods_nomenclature_item_id, 2) = '""" + sChapter + """'
ORDER BY gn.goods_nomenclature_item_id, gn.producline_suffix, gn.validity_start_date, gndp.validity_start_date;"""



# Create the Word document
document = Document()

# Add the overall document title
document.add_heading(sTitle, 0)

"""
# Add the introductory text
p = document.add_paragraph('Lorem ipsum dolor sit amet, consectetur adipiscing elit. Vivamus lacinia bibendum felis eleifend sollicitudin. Ut non pharetra justo. Phasellus sem metus, suscipit sit amet mollis sed, tempor non justo. In volutpat metus tortor, non ultrices lacus blandit sit amet. Duis leo velit, auctor vitae rhoncus in, vestibulum tempor neque. Integer ac lectus ut tellus imperdiet mattis. Vivamus vitae felis lorem. Suspendisse vel dui a leo porttitor vehicula.')

"""
#####################################################
# Section and chapter number 
document.add_heading(sSectionNumeralAndTitle.upper(), level = 1)
document.add_heading(sChapter2.upper(), level = 2)
p = document.add_paragraph(sChapterNotes2)

cur = conn.cursor()
cur.execute(sSQL)
rows = cur.fetchall()
table = document.add_table(rows = 1, cols = intColumnCount)
try:
	table.style = "List Table 3"
except:
	print ("error setting table style - investigate")
	
table.cell(0,0).width = Cm(3.0)
table.cell(0,1).width = Cm(11.0)

if (sDocumentType == "schedule"):
	table.cell(0,2).width = Cm(3.0)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Commodity code'
hdr_cells[1].text = 'Description'

if (sDocumentType == "schedule"):
	hdr_cells[2].text = 'Duty'


for row in rows:
	if row[0] != sChapter + "00000000":
		row_cells = table.add_row().cells
		sCommodityCode = row[0]
		iIndents = row[6]
		sProductLineSuffix = row[1]
		
		if (sProductLineSuffix == "80"):
			row_cells[0].text = formatCommodityCode(row[0])
			
		sDescription = ("-\t" * iIndents) + row[5]
		sDescription = sDescription.replace("|", " ")
		
		if (iIndents < 1):
			row_cells[1].paragraphs[0].add_run(sDescription).bold = True
		else:
			row_cells[1].paragraphs[0].add_run(sDescription)
			
		p = row_cells[1].paragraphs[0]
		tab_stops = p.paragraph_format.tab_stops
		for i in range(0, 12):
			tab_stop = tab_stops.add_tab_stop(Cm(0.2 * (i + 1)), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
		
		# Set indent on the cell
		row_cells[1].paragraphs[0].paragraph_format.left_indent = Cm(0.2 * (iIndents + 0))
		row_cells[1].paragraphs[0].paragraph_format.first_line_indent = Cm(-0.2 * (iIndents + 0))
		
		sDuty = ""
		
		if (sDocumentType == "schedule"):
			for rd in rdList:
				sCommodityCode2 = rd[0]
				sActive = rd[7]
				if (sCommodityCode == sCommodityCode2 and sActive == "Active"):
					if (rd[6] == ""):
						sDuty = "Free"
					else:
						sDuty = str(rd[6])
					break
		
			row_cells[2].text = sDuty

# Make the first two lines of the table repeat
set_repeat_table_header(table.rows[0])
#set_repeat_table_header(table.rows[1])

###############################################################
#Step 5 - Get the nomenclature footnotes
sSQL = """SELECT fagn.goods_nomenclature_item_id, fagn.footnote_type || fagn.footnote_id as fid, f.description
FROM footnote_association_goods_nomenclatures fagn, ml.ml_footnotes f
WHERE fagn.footnote_type = f.footnote_type_id
AND fagn.footnote_id = f.footnote_id
AND LEFT(fagn.goods_nomenclature_item_id, 2) = '""" + sChapter + """'
AND fagn.validity_start_date < CURRENT_DATE
AND (fagn.validity_end_date > CURRENT_DATE OR fagn.validity_end_date IS NULL)
ORDER BY 1, 2, 3;"""
cur = conn.cursor()
cur.execute(sSQL)
rows_footnotes = cur.fetchall()
cur.close()

listFootnotes = []
for x in rows_footnotes:
	blFound = 0
	for l in listFootnotes:
		if x[1] == l[0]:
			blFound = 1
			break
	if (blFound == 0):
		listFootnotes.append([x[1], x[2]])
	
# Add in the footnotes table at the end of the document
document.add_heading("Footnotes on chapter " + sChapter, level = 1)
tblFootnote = document.add_table(rows = 1, cols = 2)
hdr_cells = tblFootnote.rows[0].cells
hdr_cells[0].text = 'Footnote'
hdr_cells[1].text = 'Notes'

intIndex = 0
for row in listFootnotes:
	intIndex += 1
	row_cells = tblFootnote.add_row().cells
	row_cells[0].text = str(intIndex)
	#row_cells[0].text = row[0]
	row_cells[1].text = row[1]
	
# Go back over the 1st table and add in footnote references
tbl = document.tables[0]
for rw in tbl.rows:
	sCommodityCode = rw.cells[0].text

	for fn in rows_footnotes:
		#print (sCommodityCode + " : " + fn[0])
		if (sCommodityCode == formatCommodityCode(fn[0])):
			p = rw.cells[0].paragraphs[0]
			intfn2index = 0
			for fn2 in listFootnotes:
				intfn2index += 1
				if fn[1] == fn2[0]:
					print ("found match on " + sCommodityCode)
					break
				
			p.add_run(' (' + str(intfn2index) + ')').font.superscript = True
			break

	
conn.close()

###########################################################################
## Write the document	
###########################################################################

print ("here before writing")
sFilename = "output\\" + sOutTemplate.replace("XX", sChapter)
print (sFilename)
document.save(sFilename)
print ("here after writing")
	
"""
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
document.add_paragraph('first item in unordered list', style='ListBullet')
document.add_paragraph('first item in ordered list', style='ListNumber')
document.add_paragraph('Intense quote', style='IntenseQuote')
document.add_page_break()

"""
###########################################################################
###########################################################################
